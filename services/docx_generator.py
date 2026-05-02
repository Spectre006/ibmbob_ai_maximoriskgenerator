"""Word document (DOCX) report generation service."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

from utils.logger import get_logger
from models.jha_report import JHAReport
from config.constants import REPORT_TITLE, REPORT_SUBTITLE, COMPANY_NAME

logger = get_logger(__name__)


class DOCXGenerator:
    """Service for generating Word (DOCX) reports."""
    
    def __init__(self):
        """Initialize DOCX generator."""
        # IBM color palette
        self.colors = {
            'primary': RGBColor(15, 98, 254),      # IBM Blue
            'secondary': RGBColor(57, 57, 57),     # Dark gray
            'high_risk': RGBColor(218, 30, 40),    # Red
            'medium_risk': RGBColor(142, 106, 0),  # Yellow-brown
            'low_risk': RGBColor(14, 96, 39),      # Green
            'light_bg': RGBColor(244, 244, 244)    # Light gray
        }
        
        logger.info("DOCX Generator initialized")
    
    def generate_docx(self, report: JHAReport, output_path: str) -> str:
        """
        Generate Word document from JHA report data.
        
        Args:
            report: JHA report object
            output_path: Path to save DOCX file
        
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating DOCX report: {report.report_id}")
        
        try:
            # Create document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add header
            self._add_header(doc, report)
            
            # Add work order details
            self._add_work_order_section(doc, report)
            
            # Add hazards section
            self._add_hazards_section(doc, report)
            
            # Add emergency contacts
            self._add_emergency_contacts_section(doc, report)
            
            # Add footer
            self._add_footer(doc, report)
            
            # Save document
            doc.save(output_path)
            
            logger.info(f"DOCX report generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX report: {e}")
            raise Exception(f"DOCX generation failed: {str(e)}")
    
    def _add_header(self, doc: Document, report: JHAReport):
        """Add report header."""
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run(REPORT_TITLE)
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = self.colors['primary']
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(REPORT_SUBTITLE)
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = self.colors['secondary']
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company
        company = doc.add_paragraph()
        company_run = company.add_run(COMPANY_NAME)
        company_run.font.size = Pt(11)
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Report info table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('Report ID:', report.report_id),
            ('Work Order:', report.work_order_id),
            ('Generated:', report.created_at.strftime('%Y-%m-%d %H:%M:%S UTC')),
            ('Status:', report.status.value)
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Bold labels
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Spacer
    
    def _add_work_order_section(self, doc: Document, report: JHAReport):
        """Add work order details section."""
        # Section heading
        heading = doc.add_heading('Work Order Details', level=1)
        heading.runs[0].font.color.rgb = self.colors['primary']
        
        # Work order table
        wo = report.work_order
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        wo_data = [
            ('Description:', wo.get('description', 'N/A')),
            ('Location:', wo.get('location', 'N/A')),
            ('Equipment:', wo.get('equipment', 'N/A')),
            ('Priority:', wo.get('priority', 'N/A')),
            ('Status:', wo.get('status', 'N/A')),
            ('Assigned To:', wo.get('assigned_to', 'N/A'))
        ]
        
        for i, (label, value) in enumerate(wo_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Bold labels
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            # Set column widths
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.5)
        
        doc.add_paragraph()  # Spacer
    
    def _add_hazards_section(self, doc: Document, report: JHAReport):
        """Add hazards identification section."""
        # Section heading
        heading = doc.add_heading('Identified Hazards & Risk Assessment', level=1)
        heading.runs[0].font.color.rgb = self.colors['primary']
        
        if not report.hazards:
            doc.add_paragraph('No hazards identified.')
            return
        
        # Add each hazard
        for hazard in report.hazards:
            self._add_hazard_table(doc, hazard)
            doc.add_paragraph()  # Spacer between hazards
    
    def _add_hazard_table(self, doc: Document, hazard):
        """Add table for individual hazard."""
        # Hazard header
        header_para = doc.add_paragraph()
        header_para.add_run(f'Hazard #{hazard.id}').bold = True
        
        # Risk level with color
        risk_run = header_para.add_run(f' - Risk Level: {hazard.risk_level.value}')
        risk_run.bold = True
        risk_run.font.color.rgb = self._get_risk_color(hazard.risk_level.value)
        
        # Hazard details table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set background color for header
        for cell in table.rows[0].cells:
            cell._element.get_or_add_tcPr().append(
                self._create_shading_element(self.colors['light_bg'])
            )
        
        details_data = [
            ('Description:', hazard.description),
            ('Control Measures:', ', '.join(hazard.controls) if hazard.controls else 'None specified'),
            ('Required PPE:', ', '.join(hazard.ppe) if hazard.ppe else 'Standard PPE')
        ]
        
        for i, (label, value) in enumerate(details_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Bold labels
            row.cells[0].paragraphs[0].runs[0].bold = True
            
            # Set column widths
            row.cells[0].width = Inches(1.5)
            row.cells[1].width = Inches(4.5)
    
    def _add_emergency_contacts_section(self, doc: Document, report: JHAReport):
        """Add emergency contacts section."""
        # Section heading
        heading = doc.add_heading('Emergency Contacts', level=1)
        heading.runs[0].font.color.rgb = self.colors['primary']
        
        if not report.emergency_contacts:
            doc.add_paragraph('No emergency contacts specified.')
            return
        
        # Contacts table
        contacts = report.emergency_contacts
        contacts_data = []
        
        if contacts.supervisor:
            contacts_data.append(('Supervisor:', contacts.supervisor))
        if contacts.safety_officer:
            contacts_data.append(('Safety Officer:', contacts.safety_officer))
        if contacts.emergency:
            contacts_data.append(('Emergency Services:', contacts.emergency))
        
        if contacts_data:
            table = doc.add_table(rows=len(contacts_data), cols=2)
            table.style = 'Light Grid Accent 1'
            
            for i, (label, value) in enumerate(contacts_data):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                
                # Bold labels
                row.cells[0].paragraphs[0].runs[0].bold = True
                
                # Set column widths
                row.cells[0].width = Inches(2)
                row.cells[1].width = Inches(4)
        
        doc.add_paragraph()  # Spacer
    
    def _add_footer(self, doc: Document, report: JHAReport):
        """Add report footer."""
        # Additional notes
        if report.additional_notes:
            heading = doc.add_heading('Additional Notes', level=1)
            heading.runs[0].font.color.rgb = self.colors['primary']
            
            doc.add_paragraph(report.additional_notes)
            doc.add_paragraph()  # Spacer
        
        # Disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run(
            'This Job Hazard Analysis report was generated using AI-powered risk assessment. '
            'It should be reviewed by a qualified safety professional before implementation. '
            'Always follow your organization\'s safety procedures and regulations.'
        )
        disclaimer_run.italic = True
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.font.color.rgb = self.colors['secondary']
    
    def _get_risk_color(self, risk_level: str) -> RGBColor:
        """Get color for risk level."""
        risk_colors = {
            'High': self.colors['high_risk'],
            'Medium': self.colors['medium_risk'],
            'Low': self.colors['low_risk']
        }
        return risk_colors.get(risk_level, self.colors['secondary'])
    
    def _create_shading_element(self, color: RGBColor):
        """Create shading element for table cell background."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), f'{color.r:02x}{color.g:02x}{color.b:02x}')
        return shading


# Global instance
docx_generator = DOCXGenerator()

# Made with Bob
